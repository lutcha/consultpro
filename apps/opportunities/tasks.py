import logging
import os
from datetime import timedelta

from celery import shared_task
from django.utils import timezone

from .models import Opportunity, Requirement, Risk
from .scoring import score_opportunity

logger = logging.getLogger(__name__)


def _analysis_has_content(result: dict) -> bool:
    return bool(
        (result.get('summary') or '').strip()
        or result.get('requirements')
        or result.get('risks')
    )


def _coerce_text_item(item) -> str:
    if isinstance(item, dict):
        return str(item.get('description') or item.get('text') or item.get('requirement') or '').strip()
    return str(item or '').strip()


def _coerce_choice(value, allowed: set[str], default: str) -> str:
    value = str(value or '').strip().lower()
    return value if value in allowed else default


def _build_ai_extraction(result: dict, opportunity: Opportunity, service) -> dict:
    cos_analysis = result.get('cos_analysis') or {}
    return {
        'schema': 'cos_tor_analysis_v1',
        'provider': getattr(service, 'PROVIDER_NAME', 'unknown'),
        'generated_at': timezone.now().isoformat(),
        'opportunity_id': opportunity.id,
        'summary': result.get('summary', ''),
        'requirements': result.get('requirements', []),
        'risks': result.get('risks', []),
        'cos_analysis': cos_analysis,
    }


def _extract_text_from_document(opportunity: Opportunity) -> str:
    """Extract text from tor_document (DOCX or PDF) or fall back to description."""
    if not opportunity.tor_document:
        return opportunity.description or ''

    file_path = opportunity.tor_document.path
    if not os.path.exists(file_path):
        logger.warning(f"ToR file not found at {file_path}, falling back to description")
        return opportunity.description or ''

    ext = os.path.splitext(file_path)[1].lower()
    text = ''

    try:
        if ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        elif ext == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or '')
                except Exception as e:
                    logger.warning(f"Error extracting PDF page: {e}")
            text = '\n'.join(parts)
        else:
            # Try reading as plain text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception:
                text = ''
    except Exception as e:
        logger.warning(f"Error extracting text from {file_path}: {e}")

    if not text.strip():
        logger.info(f"No text extracted from ToR, falling back to description")
        text = opportunity.description or ''

    return text


def extract_text_from_uploaded_tor(file_obj) -> str:
    """Extract text from an uploaded ToR file before it leaves the API process."""
    filename = getattr(file_obj, 'name', '') or ''
    ext = os.path.splitext(filename)[1].lower()
    text = ''

    try:
        file_obj.seek(0)
        if ext == '.docx':
            from docx import Document
            doc = Document(file_obj)
            text = '\n'.join([para.text for para in doc.paragraphs])
        elif ext == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(file_obj)
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or '')
                except Exception as e:
                    logger.warning(f"Error extracting uploaded PDF page: {e}")
            text = '\n'.join(parts)
        else:
            raw = file_obj.read()
            if isinstance(raw, bytes):
                text = raw.decode('utf-8', errors='ignore')
            else:
                text = str(raw or '')
    except Exception as e:
        logger.warning(f"Error extracting text from uploaded ToR {filename}: {e}")
        text = ''
    finally:
        try:
            file_obj.seek(0)
        except Exception:
            pass

    return text


@shared_task
def analyze_tor_document(opportunity_id: int, tor_text_override: str = ''):
    """
    Analyze an Opportunity's ToR document (or description) using AI.
    Creates/updates ai_summary, Requirement, and Risk objects.
    """
    try:
        opportunity = Opportunity.objects.get(pk=opportunity_id)
    except Opportunity.DoesNotExist:
        logger.error(f"analyze_tor_document: Opportunity {opportunity_id} not found")
        return {'detail': 'Oportunidade nao encontrada.'}

    opportunity.ai_analysis_status = 'processing'
    opportunity.save(update_fields=['ai_analysis_status'])

    # AI service integration
    from apps.ai_services.services import AIServiceFactory
    service = AIServiceFactory.get_service()

    tor_text = (tor_text_override or '').strip() or _extract_text_from_document(opportunity)

    if not tor_text.strip():
        opportunity.ai_analysis_status = 'failed'
        opportunity.save(update_fields=['ai_analysis_status'])
        logger.warning(f"No text available for AI analysis on opportunity {opportunity_id}")
        return {'detail': 'Sem texto disponivel para analise.', 'opportunity_id': opportunity_id}

    # Truncate if extremely long to avoid API limits
    MAX_TEXT_LEN = 15000
    original_len = len(tor_text)
    if original_len > MAX_TEXT_LEN:
        tor_text = tor_text[:MAX_TEXT_LEN] + '\n\n[Documento truncado para analise]'

    try:
        result = service.analyze_document(tor_text)
    except Exception as e:
        logger.exception(f"AI analysis failed for opportunity {opportunity_id}: {e}")
        opportunity.ai_analysis_status = 'failed'
        opportunity.save(update_fields=['ai_analysis_status'])
        return {'detail': 'Falha na analise de IA.', 'opportunity_id': opportunity_id, 'error': str(e)}

    if not _analysis_has_content(result):
        logger.warning(
            "AI analysis returned empty output for opportunity %s using provider %s",
            opportunity_id,
            getattr(service, 'PROVIDER_NAME', 'unknown'),
        )
        opportunity.ai_analysis_status = 'failed'
        opportunity.save(update_fields=['ai_analysis_status'])
        return {
            'detail': 'Analise de IA sem resultado util.',
            'opportunity_id': opportunity_id,
        }

    summary = result.get('summary', '')
    requirements = result.get('requirements', [])
    risks = result.get('risks', [])
    ai_extraction = _build_ai_extraction(result, opportunity, service)

    # Save summary
    opportunity.ai_summary = summary
    opportunity.ai_extraction = ai_extraction

    # Create/update Requirements from AI extraction
    _sync_ai_requirements(opportunity, requirements)

    # Create/update Risks from AI extraction
    _sync_ai_risks(opportunity, risks)

    opportunity.ai_analysis_status = 'completed'
    opportunity.save(update_fields=['ai_analysis_status', 'ai_summary', 'ai_extraction'])

    logger.info(
        f"AI analysis completed for opportunity {opportunity_id}: "
        f"summary_len={len(summary)}, requirements={len(requirements)}, risks={len(risks)}"
    )

    return {
        'detail': 'Analise concluida.',
        'opportunity_id': opportunity_id,
        'requirements_created': len(requirements),
        'risks_created': len(risks),
    }


def _sync_ai_requirements(opportunity: Opportunity, requirements: list):
    """Sync AI-extracted requirements with the database."""
    if not requirements:
        return

    # Remove old AI-extracted requirements to avoid duplicates
    opportunity.requirements.filter(extracted_by_ai=True).delete()

    for item in requirements:
        req_text = _coerce_text_item(item)
        if len(req_text) < 5:
            continue
        if isinstance(item, dict):
            category = _coerce_choice(
                item.get('category'),
                {'functional', 'technical', 'institutional', 'financial'},
                'functional',
            )
            priority = _coerce_choice(
                item.get('priority'),
                {'mandatory', 'preferred', 'optional'},
                'mandatory',
            )
        else:
            category = 'functional'
            priority = 'mandatory'
            lower = req_text.lower()
            if any(k in lower for k in ['tecnico', 'technical', 'technology', 'software', 'system']):
                category = 'technical'
            elif any(k in lower for k in ['institucional', 'institutional', 'organizational', 'experience']):
                category = 'institutional'
            elif any(k in lower for k in ['financeiro', 'financial', 'budget', 'cost', 'price']):
                category = 'financial'

        Requirement.objects.create(
            opportunity=opportunity,
            category=category,
            description=req_text[:500],
            priority=priority,
            extracted_by_ai=True,
        )


def _sync_ai_risks(opportunity: Opportunity, risks: list):
    """Sync AI-extracted risks with the database."""
    if not risks:
        return

    # Remove old AI-extracted risks
    opportunity.risks.filter(identified_by_ai=True).delete()

    for item in risks:
        risk_text = _coerce_text_item(item)
        if len(risk_text) < 5:
            continue
        if isinstance(item, dict):
            severity = _coerce_choice(item.get('severity'), {'low', 'medium', 'high'}, 'medium')
            mitigation = str(item.get('mitigation') or '').strip()
        else:
            severity = 'medium'
            mitigation = ''
            lower = risk_text.lower()
            if any(k in lower for k in ['critical', 'severe', 'high impact', 'grave']):
                severity = 'high'
            elif any(k in lower for k in ['minor', 'low impact', 'small']):
                severity = 'low'

        Risk.objects.create(
            opportunity=opportunity,
            description=risk_text[:500],
            severity=severity,
            mitigation=mitigation[:500],
            identified_by_ai=True,
        )


@shared_task
def check_upcoming_deadlines():
    upcoming = Opportunity.objects.filter(
        status__in=['new', 'analyzing', 'go', 'proposal_draft', 'proposal_review'],
        deadline__gte=timezone.now(),
        deadline__lte=timezone.now() + timedelta(days=7),
    )

    # TODO: Enviar notificacoes por email/push para os utilizadores atribuidos
    count = upcoming.count()
    return {'upcoming_deadlines_count': count}


@shared_task
def enrich_and_score_opportunity(opportunity_id: int):
    try:
        opportunity = Opportunity.objects.get(pk=opportunity_id)
    except Opportunity.DoesNotExist:
        logger.error("enrich_and_score_opportunity: Opportunity %s not found", opportunity_id)
        return {'detail': 'Oportunidade nao encontrada.', 'opportunity_id': opportunity_id}

    score = score_opportunity(opportunity)
    return {
        'detail': 'Scoring concluido.',
        'opportunity_id': opportunity.id,
        'score_id': score.id,
        'overall_score': score.overall_score,
        'scoring_version': score.scoring_version,
    }
