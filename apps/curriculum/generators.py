"""Document generators for curriculum/CV adaptation.

Supports DOCX generation via python-docx and PDF generation via reportlab.
"""

import io
from pathlib import Path
from typing import Any

from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


def _set_cell_shading(cell, fill_color):
    """Set background colour for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, font_name='Calibri', font_size=11, bold=False, color=None):
    """Apply font formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def generate_cv_document(adapted_content: dict, template, output_format: str = 'docx') -> str:
    """Generate a CV document adapted for a specific template.

    Args:
        adapted_content: Dict with sections and content to include.
        template: CVTemplate instance with format rules.
        output_format: 'docx' or 'pdf'.

    Returns:
        URL/path of the generated document.
    """
    if output_format == 'pdf':
        return _generate_cv_pdf(adapted_content, template)
    return _generate_cv_docx(adapted_content, template)


def _generate_cv_docx(adapted_content: dict, template) -> str:
    """Generate a DOCX CV document."""
    doc = Document()

    # Page margins
    sections = doc.sections[0]
    sections.top_margin = Cm(2)
    sections.bottom_margin = Cm(2)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # Extract formatting rules
    format_rules = template.format_rules or []
    rules_dict = format_rules[0] if format_rules and isinstance(format_rules[0], dict) else {}

    font_name = rules_dict.get('font', 'Calibri')
    primary_color = rules_dict.get('primary_color', '#1F4E79')
    accent_color = rules_dict.get('accent_color', '#4472C4')

    # --- Header with organization branding ---
    org_name = template.organization_name or template.name
    heading = doc.add_heading(level=0)
    run = heading.add_run(org_name.upper())
    _set_run_font(run, font_name=font_name, font_size=18, bold=True, color=(31, 78, 121))
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Curriculum Vitae — {template.name}")
    _set_run_font(run, font_name=font_name, font_size=11, color=(68, 114, 196))

    doc.add_paragraph()  # Spacer

    # --- Applicant info ---
    applicant = adapted_content.get('applicant', {})
    if applicant:
        p = doc.add_paragraph()
        run = p.add_run(f"{applicant.get('name', 'Candidate')}")
        _set_run_font(run, font_name=font_name, font_size=16, bold=True)
        p.add_run('\n')
        contact_parts = []
        if applicant.get('email'):
            contact_parts.append(applicant['email'])
        if applicant.get('phone'):
            contact_parts.append(applicant['phone'])
        if applicant.get('location'):
            contact_parts.append(applicant['location'])
        if contact_parts:
            run = p.add_run('  |  '.join(contact_parts))
            _set_run_font(run, font_name=font_name, font_size=10, color=(100, 100, 100))
        doc.add_paragraph()

    # --- Professional Summary ---
    summary = adapted_content.get('summary', '')
    if summary:
        p = doc.add_paragraph()
        run = p.add_run('Professional Summary')
        _set_run_font(run, font_name=font_name, font_size=13, bold=True, color=(31, 78, 121))
        p = doc.add_paragraph(summary)
        _set_run_font(p.runs[0] if p.runs else p.add_run(), font_name=font_name, font_size=11)
        doc.add_paragraph()

    # --- Sections ---
    sections = adapted_content.get('sections', {})
    required = template.required_sections or []

    for section_def in required:
        sec_id = section_def.get('id', '')
        sec_name = section_def.get('name', sec_id)
        sec_content = sections.get(sec_id, '')

        if not sec_content:
            continue

        # Section header with bottom border effect via table
        hdr_table = doc.add_table(rows=1, cols=1)
        hdr_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_table.autofit = False
        hdr_table.columns[0].width = Inches(6)
        cell = hdr_table.rows[0].cells[0]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(sec_name.upper())
        _set_run_font(run, font_name=font_name, font_size=12, bold=True, color=(255, 255, 255))
        _set_cell_shading(cell, '1F4E79')
        cell.vertical_alignment = 1  # CENTER
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Content
        if isinstance(sec_content, list):
            for item in sec_content:
                if isinstance(item, dict):
                    _add_dict_item_to_doc(doc, item, font_name)
                else:
                    p = doc.add_paragraph(str(item), style='List Bullet')
                    if p.runs:
                        _set_run_font(p.runs[0], font_name=font_name, font_size=11)
        elif isinstance(sec_content, dict):
            _add_dict_item_to_doc(doc, sec_content, font_name)
        else:
            p = doc.add_paragraph(str(sec_content))
            if p.runs:
                _set_run_font(p.runs[0], font_name=font_name, font_size=11)

        doc.add_paragraph()  # Spacer between sections

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Document generated for {org_name} — "
        "This CV has been adapted to meet the organization's template requirements."
    )
    _set_run_font(run, font_name=font_name, font_size=9, color=(128, 128, 128))

    # Save to storage
    safe_org = ''.join(c if c.isalnum() else '_' for c in template.organization)
    file_name = f"cv_{safe_org}_{template.id}_{adapted_content.get('applicant', {}).get('name', 'candidate').replace(' ', '_')}.docx"
    output_path = f"curriculum/generated/{file_name}"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    saved_path = default_storage.save(output_path, ContentFile(buffer.read()))
    return default_storage.url(saved_path)


def _add_dict_item_to_doc(doc, item: dict, font_name: str):
    """Add a structured dict item (experience/education) to the document."""
    title = item.get('title', '')
    org = item.get('organization', '')
    date_range = item.get('date_range', '')
    description = item.get('description', '')

    # Title line
    p = doc.add_paragraph()
    run = p.add_run(title)
    _set_run_font(run, font_name=font_name, font_size=11, bold=True)
    if org:
        run = p.add_run(f"  —  {org}")
        _set_run_font(run, font_name=font_name, font_size=11, color=(68, 114, 196))
    if date_range:
        run = p.add_run(f"  ({date_range})")
        _set_run_font(run, font_name=font_name, font_size=10, color=(100, 100, 100))

    if description:
        p = doc.add_paragraph(description)
        if p.runs:
            _set_run_font(p.runs[0], font_name=font_name, font_size=10)
        p.paragraph_format.left_indent = Inches(0.25)


def _generate_cv_pdf(adapted_content: dict, template) -> str:
    """Generate a PDF CV document using reportlab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    font_name = 'Helvetica'
    font_name_bold = 'Helvetica-Bold'

    # Custom styles
    title_style = ParagraphStyle(
        'OrgTitle',
        parent=styles['Heading1'],
        fontName=font_name_bold,
        fontSize=18,
        textColor=colors.HexColor('#1F4E79'),
        alignment=1,  # CENTER
        spaceAfter=2,
    )
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=11,
        textColor=colors.HexColor('#4472C4'),
        alignment=1,
        spaceAfter=12,
    )
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName=font_name_bold,
        fontSize=12,
        textColor=colors.white,
        backColor=colors.HexColor('#1F4E79'),
        spaceAfter=6,
        spaceBefore=12,
        leftIndent=4,
        rightIndent=4,
        leading=16,
    )
    normal_style = ParagraphStyle(
        'NormalCustom',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        leading=14,
        leftIndent=20,
        bulletIndent=10,
        bulletFontName=font_name_bold,
    )

    story = []

    # --- Header ---
    org_name = template.organization_name or template.name
    story.append(Paragraph(org_name.upper(), title_style))
    story.append(Paragraph(f"Curriculum Vitae — {template.name}", subtitle_style))

    # --- Applicant ---
    applicant = adapted_content.get('applicant', {})
    if applicant:
        name = applicant.get('name', 'Candidate')
        story.append(Paragraph(f"<b>{name}</b>", ParagraphStyle(
            'Name', parent=styles['Normal'], fontName=font_name_bold, fontSize=14, spaceAfter=2,
        )))
        contact_parts = []
        if applicant.get('email'):
            contact_parts.append(applicant['email'])
        if applicant.get('phone'):
            contact_parts.append(applicant['phone'])
        if applicant.get('location'):
            contact_parts.append(applicant['location'])
        if contact_parts:
            story.append(Paragraph(
                f"<font color='#666666'>{'  |  '.join(contact_parts)}</font>",
                ParagraphStyle('Contact', parent=styles['Normal'], fontSize=9, spaceAfter=12),
            ))

    # --- Summary ---
    summary = adapted_content.get('summary', '')
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_header_style))
        story.append(Paragraph(summary, normal_style))
        story.append(Spacer(1, 6))

    # --- Sections ---
    sections = adapted_content.get('sections', {})
    required = template.required_sections or []

    for section_def in required:
        sec_id = section_def.get('id', '')
        sec_name = section_def.get('name', sec_id)
        sec_content = sections.get(sec_id, '')

        if not sec_content:
            continue

        story.append(Paragraph(sec_name.upper(), section_header_style))

        if isinstance(sec_content, list):
            for item in sec_content:
                if isinstance(item, dict):
                    story.extend(_pdf_dict_item(item, normal_style, bullet_style))
                else:
                    story.append(Paragraph(f"• {item}", bullet_style))
        elif isinstance(sec_content, dict):
            story.extend(_pdf_dict_item(sec_content, normal_style, bullet_style))
        else:
            story.append(Paragraph(str(sec_content), normal_style))

        story.append(Spacer(1, 4))

    # --- Footer ---
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        f"<font size='8' color='#888888'>"
        f"Document generated for {org_name} — "
        f"This CV has been adapted to meet the organization's template requirements."
        f"</font>",
        ParagraphStyle('Footer', parent=styles['Normal'], alignment=1),
    ))

    doc.build(story)
    buffer.seek(0)

    safe_org = ''.join(c if c.isalnum() else '_' for c in template.organization)
    file_name = f"cv_{safe_org}_{template.id}_{applicant.get('name', 'candidate').replace(' ', '_')}.pdf"
    output_path = f"curriculum/generated/{file_name}"

    saved_path = default_storage.save(output_path, ContentFile(buffer.read()))
    return default_storage.url(saved_path)


def _pdf_dict_item(item: dict, normal_style, bullet_style):
    """Convert a structured dict item to reportlab flowables."""
    elements = []
    title = item.get('title', '')
    org = item.get('organization', '')
    date_range = item.get('date_range', '')
    description = item.get('description', '')

    parts = [f"<b>{title}</b>"]
    if org:
        parts.append(f"<font color='#4472C4'> — {org}</font>")
    if date_range:
        parts.append(f"<font color='#666666'> ({date_range})</font>")

    elements.append(Paragraph(''.join(parts), normal_style))

    if description:
        elements.append(Paragraph(description, ParagraphStyle(
            'Desc', parent=normal_style, leftIndent=15, fontSize=9,
        )))
        elements.append(Spacer(1, 4))

    return elements
