import pytest
import io

from docx import Document
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APIClient

from apps.projects.models import Project
from apps.curriculum.models import Curriculum
from apps.proposals.document_generator import generate_proposal_docx, generate_proposal_pdf
from apps.proposals.models import Budget, Proposal, ProposalEvent
from apps.proposals.views import ensure_default_sections
from apps.quality_checks.models import QualityCheck
from apps.proposals.tests.factories import (
    BudgetFactory,
    OpportunityFactory,
    ProposalFactory,
    ProposalSectionFactory,
    UserFactory,
)

pytestmark = pytest.mark.django_db


@pytest.fixture
def api_client():
    return APIClient()


@pytest.fixture
def authenticated_client(api_client):
    user = UserFactory()
    api_client.force_authenticate(user=user)
    return api_client, user


class TestProposalViewSet:
    def test_list_proposals(self, authenticated_client):
        client, user = authenticated_client
        ProposalFactory(created_by=user)
        url = reverse('proposal-list')
        response = client.get(url)
        assert response.status_code == status.HTTP_200_OK
        assert len(response.data['results']) == 1

    def test_create_proposal(self, authenticated_client):
        client, user = authenticated_client
        opportunity = OpportunityFactory()
        url = reverse('proposal-list')
        data = {
            'opportunity': opportunity.id,
            'title': 'New Proposal',
            'version': 1,
            'status': 'draft',
        }
        response = client.post(url, data)
        assert response.status_code == status.HTTP_201_CREATED
        proposal = Proposal.objects.get(pk=response.data['id'])
        assert proposal.sections.count() == 7
        opportunity.refresh_from_db()
        assert opportunity.status == 'proposal_draft'

    def test_retrieve_proposal_repairs_missing_default_sections(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        assert proposal.sections.count() == 0
        url = reverse('proposal-detail', kwargs={'pk': proposal.pk})

        response = client.get(url)

        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.sections.count() == 7
        assert len(response.data['sections']) == 7

    def test_sections_action(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        section = ProposalSectionFactory(proposal=proposal)
        url = reverse('proposal-sections', kwargs={'pk': proposal.pk})
        response = client.get(url)
        assert response.status_code == status.HTTP_200_OK
        assert len(response.data) == 7
        assert any(item['id'] == section.id for item in response.data)

    def test_section_detail_get(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        section = ProposalSectionFactory(proposal=proposal)
        url = reverse(
            'proposal-section-detail',
            kwargs={'pk': proposal.pk, 'section_id': section.pk},
        )
        response = client.get(url)
        assert response.status_code == status.HTTP_200_OK
        assert response.data['id'] == section.id

    def test_section_detail_put(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        section = ProposalSectionFactory(proposal=proposal)
        url = reverse(
            'proposal-section-detail',
            kwargs={'pk': proposal.pk, 'section_id': section.pk},
        )
        response = client.put(
            url, {'content': 'Updated content'}, format='json'
        )
        assert response.status_code == status.HTTP_200_OK
        section.refresh_from_db()
        assert section.content == 'Updated content'

    def test_add_custom_section(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        ProposalSectionFactory(proposal=proposal, order=3)
        url = reverse('proposal-add-section', kwargs={'pk': proposal.pk})

        response = client.post(url, {'title': 'Plano de Comunicacao'})

        assert response.status_code == status.HTTP_201_CREATED
        assert response.data['title'] == 'Plano de Comunicacao'
        assert response.data['section_type'] == 'custom'
        assert response.data['order'] == 4

    def test_section_gap_scores_requirements_against_section_content(self, authenticated_client):
        client, user = authenticated_client
        opportunity = OpportunityFactory(
            ai_extraction={
                'cos_analysis': {
                    'methodology_blueprint': [
                        'Plano de capacitacao em turismo resiliente',
                        'Matriz de monitoria e avaliacao',
                    ],
                }
            }
        )
        opportunity.requirements.create(
            category='technical',
            priority='mandatory',
            description='Plano de capacitacao em turismo resiliente',
        )
        proposal = ProposalFactory(created_by=user, opportunity=opportunity)
        section = ProposalSectionFactory(
            proposal=proposal,
            section_type='methodology',
            title='Metodologia',
            content='<p>A metodologia inclui capacitacao em turismo resiliente.</p>',
        )
        url = reverse('proposal-section-gap', kwargs={'pk': proposal.pk})

        response = client.get(url, {'section_id': section.pk})

        assert response.status_code == status.HTTP_200_OK
        assert response.data['section']['id'] == section.id
        assert response.data['total_count'] >= 2
        assert response.data['covered_count'] >= 1
        assert response.data['score'] > 0
        assert any(item['status'] == 'missing' for item in response.data['items'])

    def test_section_gap_requires_section_id(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        url = reverse('proposal-section-gap', kwargs={'pk': proposal.pk})

        response = client.get(url)

        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_add_comment(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        section = ProposalSectionFactory(proposal=proposal)
        url = reverse('proposal-add-comment', kwargs={'pk': proposal.pk})
        response = client.post(
            url,
            {'section_id': section.id, 'text': 'Nice work'},
            format='json',
        )
        assert response.status_code == status.HTTP_201_CREATED

    def test_ai_suggest(self, authenticated_client):
        from unittest.mock import patch, MagicMock
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        section = ProposalSectionFactory(proposal=proposal)
        url = reverse('proposal-ai-suggest', kwargs={'pk': proposal.pk})
        mock_service = MagicMock()
        mock_service.generate_suggestion.return_value = 'Expanded content here.'
        with patch('apps.ai_services.services.AIServiceFactory.get_service', return_value=mock_service):
            response = client.post(
                url,
                {'section_id': section.id, 'action': 'expand'},
                format='json',
            )
        assert response.status_code == status.HTTP_201_CREATED

    def test_save_action(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        url = reverse('proposal-save', kwargs={'pk': proposal.pk})
        response = client.post(url)
        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.auto_save_status == 'saved'
        assert proposal.last_saved_at is not None

    def test_submit_action(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='ready_for_submission')
        QualityCheck.objects.create(
            proposal=proposal,
            status='completed',
            overall_score=90,
            can_submit=True,
            executed_by=user,
        )
        url = reverse('proposal-submit', kwargs={'pk': proposal.pk})
        response = client.post(url)
        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert 'ready_for_submission' in response.data['detail']

    def test_submit_action_requires_required_sections_complete(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='ready_for_submission')
        ensure_default_sections(proposal)
        for section in proposal.sections.all():
            section.is_complete = True
            section.save(update_fields=['is_complete'])
        section = proposal.sections.filter(section_type='methodology').first()
        section.is_complete = False
        section.save(update_fields=['is_complete'])

        url = reverse('proposal-submit', kwargs={'pk': proposal.pk})
        response = client.post(url)

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert 'methodology' in response.data['incomplete_required_sections']

    def test_submit_action_success_when_ready_and_complete(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='ready_for_submission')
        ensure_default_sections(proposal)
        for section in proposal.sections.all():
            section.is_complete = True
            section.save(update_fields=['is_complete'])

        url = reverse('proposal-submit', kwargs={'pk': proposal.pk})
        response = client.post(url)

        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.status == 'submitted'
        assert proposal.submitted_at is not None
        assert ProposalEvent.objects.filter(proposal=proposal, event_type='submission').exists()

    def test_submit_action_requires_passing_qc(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='ready_for_submission')
        url = reverse('proposal-submit', kwargs={'pk': proposal.pk})

        response = client.post(url)

        assert response.status_code == status.HTTP_409_CONFLICT
        proposal.refresh_from_db()
        assert proposal.status == 'ready_for_submission'

    def test_approve_for_submission_marks_ready_without_project(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='qc_check')
        QualityCheck.objects.create(
            proposal=proposal,
            status='completed',
            overall_score=90,
            can_submit=True,
            executed_by=user,
        )

        url = reverse('proposal-approve-for-submission', kwargs={'pk': proposal.pk})
        response = client.post(url)

        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.status == 'ready_for_submission'
        assert response.data['proposal_id'] == proposal.id
        assert not Project.objects.filter(proposal=proposal).exists()

    def test_approve_for_submission_is_idempotent_without_project(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='qc_check')
        QualityCheck.objects.create(
            proposal=proposal,
            status='completed',
            overall_score=90,
            can_submit=True,
            executed_by=user,
        )

        url = reverse('proposal-approve-for-submission', kwargs={'pk': proposal.pk})
        first_response = client.post(url)
        second_response = client.post(url)

        assert first_response.status_code == status.HTTP_200_OK
        assert second_response.status_code == status.HTTP_400_BAD_REQUEST
        assert Project.objects.filter(proposal=proposal).count() == 0

    def test_manager_can_approve_proposal_created_by_another_user(self, api_client):
        owner = UserFactory(role='consultant')
        manager = UserFactory(role='manager')
        proposal = ProposalFactory(created_by=owner, status='qc_check')
        QualityCheck.objects.create(
            proposal=proposal,
            status='completed',
            overall_score=90,
            can_submit=True,
            executed_by=manager,
        )
        api_client.force_authenticate(user=manager)

        url = reverse('proposal-approve-for-submission', kwargs={'pk': proposal.pk})
        response = api_client.post(url)

        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.status == 'ready_for_submission'
        assert response.data['proposal_id'] == proposal.id
        assert not Project.objects.filter(proposal=proposal).exists()

    def test_transition_status_rejects_invalid_pipeline_jump(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='draft')

        url = reverse('proposal-transition-status', kwargs={'pk': proposal.pk})
        response = client.post(url, {'status': 'submitted'}, format='json')

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        proposal.refresh_from_db()
        assert proposal.status == 'draft'

    def test_project_initiation_transition_creates_project(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user, status='contract_signed')

        url = reverse('proposal-transition-status', kwargs={'pk': proposal.pk})
        response = client.post(url, {'status': 'project_initiation'}, format='json')

        assert response.status_code == status.HTTP_200_OK
        proposal.refresh_from_db()
        assert proposal.status == 'project_initiation'
        assert response.data['project_id'] == proposal.project.id
        assert proposal.project.status == Project.Status.PLANNING
        assert proposal.project.phases.count() == 5

    def test_team_action(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        url = reverse('proposal-team', kwargs={'pk': proposal.pk})
        response = client.get(url)
        assert response.status_code == status.HTTP_200_OK

    def test_add_team_member(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        new_user = UserFactory()
        url = reverse('proposal-add-team-member', kwargs={'pk': proposal.pk})
        response = client.post(
            url,
            {
                'user_id': new_user.id,
                'role': 'Consultant',
                'hours': 10,
                'hourly_rate': 50.00,
            },
            format='json',
        )
        assert response.status_code == status.HTTP_201_CREATED

    def test_budget_get(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        BudgetFactory(proposal=proposal)
        url = reverse('proposal-budget', kwargs={'pk': proposal.pk})
        response = client.get(url)
        assert response.status_code == status.HTTP_200_OK

    def test_budget_post_create(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        url = reverse('proposal-budget', kwargs={'pk': proposal.pk})
        response = client.post(
            url,
            {'total': 5000.00, 'currency': 'EUR'},
            format='json',
        )
        assert response.status_code == status.HTTP_201_CREATED
        assert Budget.objects.filter(proposal=proposal).exists()

    def test_budget_post_update(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        BudgetFactory(proposal=proposal, total=1000.00)
        url = reverse('proposal-budget', kwargs={'pk': proposal.pk})
        response = client.post(
            url,
            {'total': 2000.00},
            format='json',
        )
        assert response.status_code == status.HTTP_200_OK
        proposal.budget.refresh_from_db()
        assert proposal.budget.total == 2000.00

    def test_team_match_returns_scores_and_missing_cv_flags(self, authenticated_client):
        client, user = authenticated_client
        proposal = ProposalFactory(created_by=user)
        member_with_cv = UserFactory()
        member_without_cv = UserFactory()
        proposal.team_members_detail.create(user=member_with_cv, role='Expert')
        proposal.team_members_detail.create(user=member_without_cv, role='Analyst')
        Curriculum.objects.create(
            user=member_with_cv,
            file_name='expert.docx',
            file_type='docx',
            file='curricula/2026/05/expert.docx',
            status='analyzed',
            extracted_data={
                'skills': ['procurement', 'governance', 'monitoring'],
                'experience': [{'title': 'Senior Consultant', 'description': 'Procurement governance'}],
                'education': [{'title': 'MSc'}],
                'languages': ['Portuguese', 'English'],
            },
        )
        proposal.opportunity.requirements.create(
            category='technical',
            description='Procurement governance monitoring',
            priority='mandatory',
        )

        url = reverse('proposal-team-match', kwargs={'pk': proposal.pk})
        response = client.get(url)

        assert response.status_code == status.HTTP_200_OK
        assert response.data['proposal_id'] == proposal.id
        assert response.data['team_size'] == 2
        assert response.data['scored_members'] == 1
        assert response.data['team_score'] > 0
        assert any(item['has_analyzed_cv'] is False for item in response.data['items'])
        assert any(item['has_analyzed_cv'] is True for item in response.data['items'])

    def test_document_exports_render_html_tables(self):
        proposal = ProposalFactory()
        ProposalSectionFactory(
            proposal=proposal,
            section_type='custom',
            title='Plano de Trabalho',
            content=(
                '<h2>Plano</h2>'
                '<table class="proposal-ai-table">'
                '<thead><tr><th>Fase</th><th>Periodo</th></tr></thead>'
                '<tbody><tr><td>Diagnostico</td><td>Mes 1</td></tr></tbody>'
                '</table>'
                '<p>Fim.</p>'
            ),
        )

        docx_bytes = generate_proposal_docx(proposal)
        document = Document(io.BytesIO(docx_bytes))

        assert len(docx_bytes) > 0
        assert any(
            cell.text == 'Fase'
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        assert len(generate_proposal_pdf(proposal)) > 0

    def test_document_exports_decode_escaped_html_tables(self):
        proposal = ProposalFactory()
        ProposalSectionFactory(
            proposal=proposal,
            section_type='custom',
            title='Plano de Trabalho',
            content=(
                '&lt;strong&gt;Plano&lt;/strong&gt;'
                '&lt;table class=&quot;proposal-ai-table&quot;&gt;'
                '&lt;thead&gt;&lt;tr&gt;&lt;th&gt;Fase&lt;/th&gt;&lt;/tr&gt;&lt;/thead&gt;'
                '&lt;tbody&gt;&lt;tr&gt;&lt;td&gt;Diagnostico&lt;/td&gt;&lt;/tr&gt;&lt;/tbody&gt;'
                '&lt;/table&gt;'
            ),
        )

        docx_bytes = generate_proposal_docx(proposal)
        document = Document(io.BytesIO(docx_bytes))

        assert len(docx_bytes) > 0
        assert any(
            cell.text == 'Fase'
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        assert len(generate_proposal_pdf(proposal)) > 0

    def test_document_exports_render_nested_html_tables(self):
        proposal = ProposalFactory()
        ProposalSectionFactory(
            proposal=proposal,
            section_type='custom',
            title='Plano de Trabalho',
            content=(
                '<p><strong>Plano</strong></p>'
                '<div>'
                '<p>Intro</p>'
                '<table class="proposal-ai-table">'
                '<thead><tr><th>Fase</th><th>Periodo</th></tr></thead>'
                '<tbody><tr><td>Diagnostico</td><td>Mes 1</td></tr></tbody>'
                '</table>'
                '</div>'
            ),
        )

        document = Document(io.BytesIO(generate_proposal_docx(proposal)))

        assert any(
            cell.text == 'Diagnostico'
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

    def test_document_exports_rebuild_linearized_workplan_table(self):
        proposal = ProposalFactory()
        ProposalSectionFactory(
            proposal=proposal,
            section_type='workplan',
            title='Plano de Trabalho',
            content=(
                '<p>Fase</p><p>Periodo</p><p>Atividades Principais</p><p>Entregaveis</p>'
                '<p>1. Diagnostico Inicial</p><p>Mes 1</p>'
                '<p>Reunioes individuais</p><p>Diagnostico inicial</p>'
                '<p>2. Planeamento</p><p>Mes 1</p>'
                '<p>Definicao de metas</p><p>Plano individual</p>'
            ),
        )

        document = Document(io.BytesIO(generate_proposal_docx(proposal)))

        assert any(
            [cell.text for cell in row.cells] == [
                '1. Diagnostico Inicial',
                'Mes 1',
                'Reunioes individuais',
                'Diagnostico inicial',
            ]
            for table in document.tables
            for row in table.rows
        )
