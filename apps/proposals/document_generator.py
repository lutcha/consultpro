#!/usr/bin/env python
"""
Document Generator for Proposals.
Generates professional Word (.docx) and PDF (.pdf) documents from proposal data.
Supports placeholders for company and client logos.
"""

import io
import os
from datetime import date
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from apps.proposals.models import Proposal


def _set_cell_shading(cell, fill_color: str):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading."""
    paragraph = doc.add_heading(text, level=level)
    run = paragraph.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)  # Dark blue
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    return paragraph


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False, alignment=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p


def _html_text(node) -> str:
    return node.get_text(' ', strip=True) if hasattr(node, 'get_text') else str(node).strip()


def _iter_html_blocks(html: str):
    soup = BeautifulSoup(html or '', 'html.parser')
    source = soup.body if soup.body else soup
    for child in source.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                yield 'p', text
            continue
        if isinstance(child, Tag):
            yield child.name, child


def _add_html_table_docx(doc: Document, table_node: Tag):
    rows = table_node.find_all('tr')
    if not rows:
        return

    first_cells = rows[0].find_all(['th', 'td'])
    column_count = max(len(first_cells), 1)
    table = doc.add_table(rows=1, cols=column_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row_node in enumerate(rows):
        cells = row_node.find_all(['th', 'td'])
        row_cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for cell_index in range(column_count):
            row_cells[cell_index].text = _html_text(cells[cell_index]) if cell_index < len(cells) else ''
            if row_index == 0:
                _set_cell_shading(row_cells[cell_index], 'E2E8F0')
                for paragraph in row_cells[cell_index].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)


def _add_html_content_docx(doc: Document, html: str):
    for tag_name, node in _iter_html_blocks(html):
        if tag_name == 'table' and isinstance(node, Tag):
            _add_html_table_docx(doc, node)
            doc.add_paragraph()
        elif tag_name in {'h1', 'h2', 'h3'} and isinstance(node, Tag):
            level = 2 if tag_name == 'h1' else 3
            _add_heading(doc, _html_text(node), level=level)
        elif tag_name in {'ul', 'ol'} and isinstance(node, Tag):
            style = 'List Bullet' if tag_name == 'ul' else 'List Number'
            for item in node.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(style=style)
                paragraph.add_run(_html_text(item))
        else:
            text = _html_text(node)
            if text:
                _add_paragraph(doc, text)


def _add_html_content_pdf(story: list, html: str, normal_style, heading2_style):
    for tag_name, node in _iter_html_blocks(html):
        if tag_name == 'table' and isinstance(node, Tag):
            rows = []
            for row_node in node.find_all('tr'):
                cells = row_node.find_all(['th', 'td'])
                rows.append([Paragraph(_html_text(cell), normal_style) for cell in cells])
            if rows:
                column_count = max(len(row) for row in rows)
                for row in rows:
                    row.extend([''] * (column_count - len(row)))
                table = Table(rows, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E2E8F0')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1A365D')),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CBD5E1')),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.25*cm))
        elif tag_name in {'h1', 'h2', 'h3'} and isinstance(node, Tag):
            story.append(Paragraph(_html_text(node), heading2_style))
        elif tag_name in {'ul', 'ol'} and isinstance(node, Tag):
            for item in node.find_all('li', recursive=False):
                bullet = '•' if tag_name == 'ul' else '•'
                story.append(Paragraph(f"{bullet} {_html_text(item)}", normal_style))
        else:
            text = _html_text(node)
            if text:
                story.append(Paragraph(text, normal_style))


def generate_proposal_docx(proposal: Proposal, output_path: Optional[str] = None) -> bytes:
    """
    Generate a Word document (.docx) for a proposal.
    
    Returns the document bytes if output_path is None.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # --- HEADER WITH LOGO PLACEHOLDERS ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(3)
    
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    
    # Left: Proponent company logo placeholder
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[LOGO EMPRESA PROPONENTE]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    p.add_run(f"\n{proposal.created_by.get_full_name() or proposal.created_by.username if proposal.created_by else 'ConsultPro'}")
    
    # Right: Client logo placeholder
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("[LOGO CLIENTE]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    if proposal.opportunity:
        p.add_run(f"\n{proposal.opportunity.client}")
    
    doc.add_paragraph()  # Spacer
    
    # --- TITLE PAGE ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(proposal.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    doc.add_paragraph()
    
    # Opportunity reference
    if proposal.opportunity:
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_run = ref_para.add_run(f"Referência: {proposal.opportunity.reference_number or 'N/A'}")
        ref_run.font.size = Pt(12)
        ref_run.italic = True
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Data: {date.today().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # --- CONSORTIUM INFO (if applicable) ---
    consortium_members = getattr(proposal, 'consortium_members', None)
    if consortium_members:
        _add_heading(doc, "Membros do Consórcio", level=1)
        for member in consortium_members:
            _add_paragraph(doc, f"• {member}", bold=False)
        doc.add_paragraph()
    
    # --- TABLE OF CONTENTS PLACEHOLDER ---
    _add_heading(doc, "Índice", level=1)
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("[Índice gerado automaticamente - atualizar no Word]")
    toc_run.font.italic = True
    toc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_page_break()
    
    # --- PROPOSAL SECTIONS ---
    sections = proposal.sections.all().order_by('order')
    for section in sections:
        _add_heading(doc, section.title, level=1)
        
        if section.content:
            _add_html_content_docx(doc, section.content)
        else:
            p = doc.add_paragraph()
            run = p.add_run("[Conteúdo em desenvolvimento...]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        
        doc.add_paragraph()  # Spacer between sections
    
    # --- TEAM SECTION ---
    team_members = proposal.team_members_detail.all()
    if team_members:
        doc.add_page_break()
        _add_heading(doc, "Equipa da Proposta", level=1)
        
        team_table = doc.add_table(rows=1, cols=4)
        team_table.style = 'Table Grid'
        team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = team_table.rows[0].cells
        headers = ['Nome', 'Função', 'Horas', 'Taxa Horária']
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            _set_cell_shading(hdr_cells[i], '1A365D')
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        
        # Data rows
        for member in team_members:
            row_cells = team_table.add_row().cells
            user = member.user
            row_cells[0].text = f"{user.first_name} {user.last_name}".strip() or user.username
            row_cells[1].text = member.role
            row_cells[2].text = str(member.hours)
            row_cells[3].text = f"{member.hourly_rate} EUR" if member.hourly_rate else "N/A"
    
    # --- BUDGET SECTION ---
    budget = getattr(proposal, 'budget', None)
    if budget:
        doc.add_page_break()
        _add_heading(doc, "Orçamento", level=1)
        
        _add_paragraph(doc, f"Total: {budget.total} {budget.currency}", bold=True)
        
        if hasattr(budget, 'items') and budget.items.exists():
            budget_table = doc.add_table(rows=1, cols=3)
            budget_table.style = 'Table Grid'
            
            hdr_cells = budget_table.rows[0].cells
            hdr_headers = ['Categoria', 'Descrição', 'Valor']
            for i, h in enumerate(hdr_headers):
                hdr_cells[i].text = h
                _set_cell_shading(hdr_cells[i], '1A365D')
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            
            for item in budget.items.all():
                row_cells = budget_table.add_row().cells
                row_cells[0].text = item.category
                row_cells[1].text = item.description or ""
                row_cells[2].text = str(item.amount)
    
    # --- FOOTER ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"\nDocumento gerado automaticamente pela plataforma ConsultPro — {date.today().year}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Save to bytes or file
    if output_path:
        doc.save(output_path)
        return b''
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_proposal_pdf(proposal: Proposal, output_path: Optional[str] = None) -> bytes:
    """
    Generate a PDF document for a proposal using ReportLab.
    
    Returns the PDF bytes if output_path is None.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm,
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#1A365D'),
        spaceAfter=30,
        alignment=1,  # Center
    )
    
    heading1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1A365D'),
        spaceAfter=12,
        spaceBefore=12,
    )
    
    heading2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.HexColor('#1A365D'),
        spaceAfter=10,
        spaceBefore=10,
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )
    
    placeholder_style = ParagraphStyle(
        'Placeholder',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        fontName='Helvetica-Oblique',
    )
    
    story = []
    
    # --- HEADER WITH LOGO PLACEHOLDERS ---
    header_data = [
        [
            Paragraph("<i>[LOGO EMPRESA<br/>PROPONENTE]</i>", placeholder_style),
            Paragraph("<i>[LOGO CLIENTE]</i>", placeholder_style),
        ]
    ]
    header_table = Table(header_data, colWidths=[8*cm, 8*cm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 1*cm))
    
    # --- TITLE ---
    story.append(Paragraph(proposal.title, title_style))
    story.append(Spacer(1, 0.5*cm))
    
    # Reference info
    if proposal.opportunity:
        story.append(Paragraph(
            f"<i>Referência: {proposal.opportunity.reference_number or 'N/A'}</i>",
            normal_style
        ))
        story.append(Paragraph(
            f"<i>Data: {date.today().strftime('%d/%m/%Y')}</i>",
            normal_style
        ))
        story.append(Paragraph(
            f"<i>Cliente: {proposal.opportunity.client}</i>",
            normal_style
        ))
    story.append(Spacer(1, 1*cm))
    
    # --- CONSORTIUM INFO ---
    consortium_members = getattr(proposal, 'consortium_members', None)
    if consortium_members:
        story.append(Paragraph("Membros do Consórcio", heading1_style))
        for member in consortium_members:
            story.append(Paragraph(f"• {member}", normal_style))
        story.append(Spacer(1, 0.5*cm))
    
    story.append(PageBreak())
    
    # --- PROPOSAL SECTIONS ---
    sections = proposal.sections.all().order_by('order')
    for section in sections:
        story.append(Paragraph(section.title, heading1_style))
        
        if section.content:
            _add_html_content_pdf(story, section.content, normal_style, heading2_style)
        else:
            story.append(Paragraph(
                "<i>[Conteúdo em desenvolvimento...]</i>",
                placeholder_style
            ))
        
        story.append(Spacer(1, 0.5*cm))
    
    # --- TEAM SECTION ---
    team_members = proposal.team_members_detail.all()
    if team_members:
        story.append(PageBreak())
        story.append(Paragraph("Equipa da Proposta", heading1_style))
        story.append(Spacer(1, 0.3*cm))
        
        team_data = [['Nome', 'Função', 'Horas', 'Taxa Horária']]
        for member in team_members:
            user = member.user
            name = f"{user.first_name} {user.last_name}".strip() or user.username
            team_data.append([
                name,
                member.role,
                str(member.hours),
                f"{member.hourly_rate} EUR" if member.hourly_rate else "N/A"
            ])
        
        team_table = Table(team_data, colWidths=[6*cm, 4*cm, 3*cm, 3*cm])
        team_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A365D')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7FAFC')]),
        ]))
        story.append(team_table)
    
    # --- BUDGET SECTION ---
    budget = getattr(proposal, 'budget', None)
    if budget:
        story.append(PageBreak())
        story.append(Paragraph("Orçamento", heading1_style))
        story.append(Paragraph(
            f"<b>Total: {budget.total} {budget.currency}</b>",
            normal_style
        ))
        story.append(Spacer(1, 0.3*cm))
        
        if hasattr(budget, 'items') and budget.items.exists():
            budget_data = [['Categoria', 'Descrição', 'Valor']]
            for item in budget.items.all():
                budget_data.append([
                    item.category,
                    item.description or "",
                    str(item.amount)
                ])
            
            budget_table = Table(budget_data, colWidths=[5*cm, 8*cm, 3*cm])
            budget_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A365D')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7FAFC')]),
            ]))
            story.append(budget_table)
    
    # --- FOOTER ---
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(
        f"<i>Documento gerado automaticamente pela plataforma ConsultPro — {date.today().year}</i>",
        placeholder_style
    ))
    
    doc.build(story)
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(buffer.getvalue())
    
    buffer.seek(0)
    return buffer.getvalue()
